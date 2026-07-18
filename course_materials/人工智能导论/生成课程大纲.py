from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

document = Document()

style = document.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

title = document.add_heading('人工智能导论课程大纲', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(22)
title.runs[0].font.bold = True

document.add_paragraph()

info_table = document.add_table(rows=4, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

info_data = [
    ['课程名称', '人工智能导论'],
    ['课程代码', 'CSAI101'],
    ['学分', '3'],
    ['学时', '48'],
]

for i, row_data in enumerate(info_data):
    for j, cell_data in enumerate(row_data):
        cell = info_table.cell(i, j)
        cell.text = cell_data
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.name = '宋体'
        cell.paragraphs[0].runs[0].font.size = Pt(12)

document.add_paragraph()

document.add_heading('一、课程概述', level=1)

document.add_paragraph('''本课程是人工智能领域的入门课程，旨在向学生介绍人工智能的基本概念、核心技术和应用领域。通过学习本课程，学生将了解人工智能的发展历程、掌握机器学习和深度学习的基础理论，熟悉自然语言处理、计算机视觉等核心技术，并具备初步的AI实践能力。''')

document.add_heading('二、课程目标', level=1)

objectives = [
    '1. 理解人工智能的基本概念和发展历程',
    '2. 掌握机器学习的基本原理和方法',
    '3. 理解深度学习的核心架构和算法',
    '4. 熟悉自然语言处理和计算机视觉的基本技术',
    '5. 了解强化学习和知识图谱的基本概念',
    '6. 具备使用Python进行AI开发的基本能力',
    '7. 理解AI伦理和安全问题',
]

for obj in objectives:
    p = document.add_paragraph(obj)
    p.runs[0].font.size = Pt(12)

document.add_heading('三、教学内容与学时分配', level=1)

content_table = document.add_table(rows=9, cols=4)
content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
content_table.style = 'Table Grid'

header_row = content_table.rows[0]
headers = ['章节', '教学内容', '学时', '教学方式']
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)

content_data = [
    ['第1章', '人工智能概述', '6', '课堂讲授+讨论'],
    ['第2章', '机器学习基础', '8', '课堂讲授+实验'],
    ['第3章', '深度学习基础', '8', '课堂讲授+实验'],
    ['第4章', '自然语言处理', '6', '课堂讲授+实验'],
    ['第5章', '计算机视觉', '6', '课堂讲授+实验'],
    ['第6章', '强化学习', '6', '课堂讲授+讨论'],
    ['第7章', '知识图谱', '4', '课堂讲授'],
    ['第8章', 'AI实践与应用', '4', '项目实践'],
]

for i, row_data in enumerate(content_data):
    for j, cell_data in enumerate(row_data):
        cell = content_table.rows[i+1].cells[j]
        cell.text = cell_data
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(11)

document.add_paragraph()

document.add_heading('四、教学方法', level=1)

methods = [
    '1. 课堂讲授：系统讲解AI基础理论和核心技术',
    '2. 案例分析：结合实际案例理解AI应用',
    '3. 实验教学：通过编程实践加深理解',
    '4. 课堂讨论：促进学生思考和交流',
    '5. 项目实践：综合运用所学知识完成AI项目',
]

for method in methods:
    p = document.add_paragraph(method)
    p.runs[0].font.size = Pt(12)

document.add_heading('五、教材与参考资料', level=1)

document.add_paragraph('【主教材】')
document.add_paragraph('《人工智能导论》，作者：XXX，出版社：XXX，出版年份：2024')

document.add_paragraph()
document.add_paragraph('【参考资料】')
refs = [
    '1. 《人工智能：一种现代方法》，Stuart Russell等，清华大学出版社',
    '2. 《深度学习》，Ian Goodfellow等，人民邮电出版社',
    '3. 《机器学习》，周志华，清华大学出版社',
    '4. 《自然语言处理入门》，何晗，人民邮电出版社',
    '5. 《计算机视觉：算法与应用》，Richard Szeliski，清华大学出版社',
]

for ref in refs:
    p = document.add_paragraph(ref)
    p.runs[0].font.size = Pt(11)

document.add_heading('六、考核方式', level=1)

assessment_table = document.add_table(rows=5, cols=3)
assessment_table.alignment = WD_TABLE_ALIGNMENT.CENTER
assessment_table.style = 'Table Grid'

header_row = assessment_table.rows[0]
headers = ['考核项目', '占比', '说明']
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)

assessment_data = [
    ['平时成绩', '30%', '考勤、作业、课堂表现'],
    ['实验成绩', '20%', '实验报告、编程实践'],
    ['期中考核', '20%', '知识测验'],
    ['期末考核', '30%', '综合考试'],
]

for i, row_data in enumerate(assessment_data):
    for j, cell_data in enumerate(row_data):
        cell = assessment_table.rows[i+1].cells[j]
        cell.text = cell_data
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(11)

document.add_paragraph()

document.add_heading('七、课程要求', level=1)

requirements = [
    '1. 具备Python编程基础',
    '2. 熟悉线性代数和概率论知识',
    '3. 按时完成作业和实验',
    '4. 积极参与课堂讨论',
    '5. 独立完成期末项目',
]

for req in requirements:
    p = document.add_paragraph(req)
    p.runs[0].font.size = Pt(12)

document.add_paragraph()

document.add_heading('八、教学进度安排', level=1)

schedule_table = document.add_table(rows=18, cols=3)
schedule_table.alignment = WD_TABLE_ALIGNMENT.CENTER
schedule_table.style = 'Table Grid'

header_row = schedule_table.rows[0]
headers = ['周次', '教学内容', '作业/实验']
for i, header in enumerate(headers):
    cell = header_row.cells[i]
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)

schedule_data = [
    ['1', '第1章：人工智能概述', '思考题'],
    ['2', '第1章：AI分类与应用', '案例分析'],
    ['3', '第2章：机器学习基础', '实验1：Python环境搭建'],
    ['4', '第2章：监督学习', '实验2：线性回归'],
    ['5', '第2章：无监督学习', '实验3：K-Means聚类'],
    ['6', '第2章：评估指标与过拟合', '作业'],
    ['7', '第3章：神经网络基础', '实验4：感知机'],
    ['8', '第3章：CNN', '实验5：图像分类'],
    ['9', '第3章：RNN与LSTM', '实验6：文本生成'],
    ['10', '期中考试', '复习'],
    ['11', '第4章：自然语言处理', '实验7：词向量'],
    ['12', '第4章：预训练模型', '实验8：BERT微调'],
    ['13', '第5章：计算机视觉', '实验9：目标检测'],
    ['14', '第6章：强化学习', '实验10：Q-Learning'],
    ['15', '第7章：知识图谱', '作业'],
    ['16', '第8章：AI实践与应用', '项目实践'],
    ['17', '课程总结与答辩', '项目展示'],
]

for i, row_data in enumerate(schedule_data):
    for j, cell_data in enumerate(row_data):
        cell = schedule_table.rows[i+1].cells[j]
        cell.text = cell_data
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(10)

document.save('课程大纲.docx')
print('课程大纲.docx 已生成')