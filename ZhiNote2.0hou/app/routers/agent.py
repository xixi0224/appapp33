import json
import os
from fastapi import APIRouter, HTTPException, Body, UploadFile, File, Depends, Cookie
from fastapi.responses import FileResponse, JSONResponse
from typing import List, Dict, Any, Optional
from sqlalchemy.orm import Session

from app.services.agent_service import (
    ProfileAgent, PlanningAgent, TeachingAgent, QuestionAgent, 
    ResourceAgent, EvaluationAgent, get_agent_status,
    generate_recommendations, update_plan_based_on_evaluation,
    complete_learning_loop, extract_topics_from_document, ocr_analyze_image,
    automated_learning_loop
)
from app.database import UploadedDocument
from app.services.db_service import (
    get_db, get_or_create_session, save_profile, get_profile,
    save_plan, get_plan, update_plan_status, save_mistake, get_mistakes,
    save_record, get_records, save_resource, get_resources,
    save_message, get_messages, clear_messages, save_document, get_documents, get_document_content,
    delete_document, mark_mistake_reviewed, mark_mistake_unreviewed, delete_mistake, clear_mistakes,
    save_question, get_questions, get_question_by_id, update_mistake, save_mistake_full,
    get_mistakes_full, get_mistakes_by_subject, get_mistakes_stats, get_daily_trend,
    save_knowledge_point, get_knowledge_points_by_document, get_knowledge_points,
    delete_knowledge_points_by_document, update_knowledge_point, get_knowledge_stats,
    clear_all_knowledge_points
)

router = APIRouter()

@router.post("/agent/session")
async def create_or_get_session(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        response = JSONResponse(content={"code": 0, "data": {"session_key": db_session.session_key}})
        response.set_cookie(
            key="session_key",
            value=db_session.session_key,
            httponly=True,
            max_age=86400 * 30,
            path="/"
        )
        return response
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def detect_intent(message):
    message_lower = message.lower()
    if any(keyword in message_lower for keyword in ['讲解', '解释', '概念', '是什么', '什么是', '原理', '定义']):
        return 'explain'
    if any(keyword in message_lower for keyword in ['计划', '路径', '规划', '路线', '安排']):
        return 'plan'
    if any(keyword in message_lower for keyword in ['题目', '习题', '练习', '做题', '测试']):
        return 'question'
    if any(keyword in message_lower for keyword in ['资源', '资料', '文档', '推荐']):
        return 'resource'
    if any(keyword in message_lower for keyword in ['评估', '分析', '错题', '正确率']):
        return 'evaluate'
    return 'general'

@router.post("/agent/chat")
async def chat(message: str = Body(..., embed=True), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        session_id = db_session.id
        
        save_message(db, session_id, 'user', message)
        
        history = get_messages(db, session_id, limit=10)
        history_str = "\n".join([f"{m['role']}: {m['content']}" for m in history])
        
        existing_profile = get_profile(db, session_id)
        
        if existing_profile:
            profile = existing_profile.profile_json if existing_profile.profile_json else {}
        else:
            profile = {}
        
        profile_str = json.dumps(profile, ensure_ascii=False)
        
        from app.services.agent_service import call_llm
        
        prompt = f"""
你是一位聪明、专业的学习助手，像豆包一样与用户自然对话。

学生画像：{profile_str}

对话历史：
{history_str}

当前用户问：{message}

请根据以上信息，执行以下操作：
1. 如果学生画像为空或信息不足，通过自然对话主动询问学生的专业、学习目标、知识基础等信息，不要生硬地问问题，要自然融入对话。
2. 如果学生画像已有足够信息，根据学生的知识水平和学习目标，提供针对性的回答。
3. 如果用户问的是概念讲解，要通俗易懂，多用生活类比。
4. 如果用户问的是具体问题，要给出清晰准确的解答。
5. 如果用户需要学习规划或资源推荐，可以主动建议查看学习路径或资源中心。
6. 回答要用中文，语气友好亲切，像朋友聊天一样自然。

请直接给出回答，不要输出JSON格式。
"""
        
        system_prompt = "你是一位聪明、友好、专业的学习助手，与用户自然对话，帮助用户解决学习问题。"
        response = call_llm(prompt, system_prompt)
        
        if not response:
            response = "抱歉，AI服务暂时不可用，请稍后再试。"
        
        save_message(db, session_id, 'assistant', response)
        save_record(db, session_id, {'type': 'chat', 'topic': message[:50], 'detail': response[:100]})
        
        if len(history) >= 3 and not profile:
            profile_agent = ProfileAgent()
            profile = profile_agent.analyze(history_str)
            save_profile(db, session_id, profile)
            
            planning_agent = PlanningAgent()
            plan = planning_agent.generate_plan(profile)
            save_plan(db, session_id, plan)
            
            response += "\n\n📋 已为你自动生成个性化学习路径，请查看「学习路径」页面。"
            
        elif profile and len(history) % 5 == 0:
            profile_agent = ProfileAgent()
            updated_profile = profile_agent.analyze(history_str)
            if updated_profile:
                save_profile(db, session_id, updated_profile)
        
        return {"code": 0, "data": {"response": response, "profile": profile}}
    except Exception as e:
        print(f"Chat error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/chat/history")
async def get_chat_history(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        messages = get_messages(db, db_session.id)
        return {"code": 0, "data": messages}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/chat/clear")
async def clear_chat_history(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        print(f"Clear history called with session_key: {session_key}")
        db_session = get_or_create_session(db, session_key)
        print(f"Session ID: {db_session.id}")
        clear_messages(db, db_session.id)
        return {"code": 0, "data": {"message": "历史记录已清空"}}
    except Exception as e:
        print(f"Clear history error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/profile")
async def update_profile(profile_data: dict = Body(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        save_profile(db, db_session.id, profile_data)
        return {"code": 0, "data": profile_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/profile")
async def get_profile_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        profile = get_profile(db, db_session.id)
        if profile:
            profile_data = profile.profile_json
            profile_data['update_history'] = profile.update_history or []
            profile_data['created_at'] = profile.created_at.isoformat() if profile.created_at else ''
            profile_data['updated_at'] = profile.updated_at.isoformat() if profile.updated_at else ''
            return {"code": 0, "data": profile_data}
        return {"code": 0, "data": {}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/plan")
async def get_plan_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        print(f"GET /agent/plan called with session_key: {session_key}")
        db_session = get_or_create_session(db, session_key)
        print(f"Session ID: {db_session.id}")
        
        try:
            plan = get_plan(db, db_session.id)
            print(f"Plan data count: {len(plan)}")
            return {"code": 0, "data": plan}
        except Exception as plan_error:
            print(f"get_plan error: {plan_error}")
            import traceback
            traceback.print_exc()
            return {"code": 0, "data": []}
    except Exception as e:
        print(f"GET /agent/plan error: {e}")
        import traceback
        traceback.print_exc()
        return {"code": 0, "data": []}

@router.post("/agent/plan")
async def create_plan(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        print("POST /agent/plan called")
        db_session = get_or_create_session(db, session_key)
        print(f"Session ID: {db_session.id}")
        
        profile = get_profile(db, db_session.id)
        if not profile:
            print("No profile found, using default profile")
            profile_data = {
                'major': '',
                'knowledge_base': 50,
                'learning_goal': 50,
                'learning_speed': 50,
                'error_patterns': 50,
                'cognitive_style': 50,
                'interest_direction': 50,
                'weak_points': [],
                'ai_summary': ''
            }
        else:
            profile_data = profile.profile_json
            print(f"Profile data loaded")
        
        try:
            docs = db.query(UploadedDocument).filter(UploadedDocument.session_id == db_session.id).all()
            knowledge_points = []
            for doc in docs:
                if doc.status == 'processed' and doc.topics_count > 0:
                    points = get_knowledge_points_by_document(db, doc.id)
                    knowledge_points.extend([{
                        'id': p.id,
                        'name': p.name,
                        'subject': p.subject or '',
                        'chapter': p.chapter or '',
                        'definition': p.definition or '',
                        'difficulty': p.difficulty or 'medium',
                        'is_key': p.is_key,
                        'is_exam_point': p.is_exam_point,
                        'prerequisites': p.prerequisites or []
                    } for p in points])
            print(f"Knowledge points count: {len(knowledge_points)}")
        except Exception as doc_error:
            print(f"Error getting documents: {doc_error}")
            knowledge_points = []
        
        planning_agent = PlanningAgent()
        plan = planning_agent.generate_plan(profile_data, knowledge_points)
        
        print(f"Generated plan count: {len(plan)}")
        
        if not plan or len(plan) == 0:
            raise HTTPException(status_code=500, detail="AI生成学习路径失败，请重试")
        
        save_plan(db, db_session.id, plan)
        
        return {"code": 0, "data": plan}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Create plan error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"生成学习路径失败: {str(e)}")

@router.put("/agent/plan/status")
async def update_plan_api(step: int = Body(..., embed=True), status: str = Body(..., embed=True), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        success = update_plan_status(db, db_session.id, step, status)
        if success:
            save_record(db, db_session.id, {'type': 'plan', 'topic': f'Step {step}', 'detail': status})
            return {"code": 0, "data": {"step": step, "status": status}}
        else:
            raise HTTPException(status_code=404, detail="步骤不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/agent/plan/update")
async def adjust_plan(evaluation_result: dict = Body(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        profile = get_profile(db, db_session.id)
        plan = get_plan(db, db_session.id)
        
        if not profile or not plan:
            raise HTTPException(status_code=400, detail="请先创建学习画像和计划")
        
        updated_plan = update_plan_based_on_evaluation(evaluation_result, profile.profile_json, plan)
        save_plan(db, db_session.id, updated_plan)
        
        return {"code": 0, "data": {"plan": updated_plan}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/resource")
async def generate_resource(topic: str = Body(..., embed=True), resource_type: str = Body(..., embed=True), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        print(f"Generate resource called: topic={topic}, type={resource_type}, session_key={session_key}")
        db_session = get_or_create_session(db, session_key)
        print(f"Session ID: {db_session.id}")
        
        content = ""
        
        if resource_type == 'document':
            resource_agent = ResourceAgent()
            content = resource_agent.generate_document(topic)
        elif resource_type == 'mindmap':
            resource_agent = ResourceAgent()
            content = resource_agent.generate_mindmap(topic)
        elif resource_type == 'ppt':
            resource_agent = ResourceAgent()
            content = resource_agent.generate_ppt(topic)
        elif resource_type == 'code':
            resource_agent = ResourceAgent()
            content = resource_agent.generate_code(topic)
        elif resource_type == 'reading':
            resource_agent = ResourceAgent()
            content = resource_agent.generate_reading(topic)
        elif resource_type == 'video_script':
            resource_agent = ResourceAgent()
            content = resource_agent.generate_video_script(topic)
        elif resource_type == 'questions':
            question_agent = QuestionAgent()
            questions = question_agent.generate_questions(topic, count=5)
            content = json.dumps(questions, ensure_ascii=False, indent=2)
        else:
            resource_agent = ResourceAgent()
            content = resource_agent.generate_document(topic)
        
        print(f"Generated content length: {len(str(content))}")
        
        if not content:
            raise Exception("AI生成内容为空，请检查API密钥配置")
        
        save_resource(db, db_session.id, {'type': resource_type, 'title': topic, 'content': str(content), 'topic': topic})
        
        return {"code": 0, "data": {"type": resource_type, "content": content}}
    except Exception as e:
        print(f"Generate resource error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/resources")
async def get_resources_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        resources = get_resources(db, db_session.id)
        return {"code": 0, "data": resources}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/questions")
async def generate_questions(topic: str = Body(..., embed=True), difficulty: str = Body(default="medium", embed=True), count: int = Body(default=5, embed=True), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        question_agent = QuestionAgent()
        questions = question_agent.generate_questions(topic, difficulty, count)
        
        for q in questions:
            save_question(db, db_session.id, {
                'subject': q.get('subject', '计算机'),
                'topic': q.get('topic', topic),
                'question': q.get('question', ''),
                'question_type': q.get('question_type', 'choice'),
                'options': q.get('options', []),
                'correct_answer': q.get('correct_answer', ''),
                'analysis': q.get('analysis', ''),
                'difficulty': q.get('difficulty', difficulty),
                'error_tags': q.get('error_types', []),
                'related_document_id': 0,
                'generated_from': 'agent'
            })
        
        return {"code": 0, "data": questions}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/questions")
async def get_questions_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        questions = get_questions(db, db_session.id)
        return {"code": 0, "data": questions}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/question/{question_id}")
async def get_single_question(question_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        question = get_question_by_id(db, question_id)
        if not question:
            raise HTTPException(status_code=404, detail="题目不存在")
        return {"code": 0, "data": question}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/evaluate")
async def evaluate(answers: List[Dict] = Body(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        profile = get_profile(db, db_session.id)
        
        if not profile:
            raise HTTPException(status_code=400, detail="请先创建学习画像")
        
        evaluation_agent = EvaluationAgent()
        result = evaluation_agent.evaluate(answers, profile.profile_json)
        
        for answer in answers:
            is_correct = answer.get('user_answer') == answer.get('correct_answer')
            
            save_record(db, db_session.id, {
                'type': 'question',
                'topic': answer.get('topic', ''),
                'detail': answer.get('question', '')[:100],
                'score': 100 if is_correct else 0
            })
            
            if not is_correct:
                save_mistake_full(db, db_session.id, {
                    'question_id': answer.get('question_id', 0),
                    'question': answer.get('question', ''),
                    'question_type': answer.get('question_type', 'choice'),
                    'user_answer': answer.get('user_answer', ''),
                    'correct_answer': answer.get('correct_answer', ''),
                    'analysis': answer.get('explanation', ''),
                    'error_tags': answer.get('error_types', []),
                    'error_detail': answer.get('error_detail', ''),
                    'related_knowledge': answer.get('topic', ''),
                    'subject': answer.get('subject', '计算机'),
                    'topic': answer.get('topic', ''),
                    'difficulty': answer.get('difficulty', 'medium'),
                    'related_document_id': answer.get('related_document_id', 0),
                    'source_type': 'agent'
                })
        
        if result.get('updated_profile'):
            profile_data = profile.profile_json
            profile_data['dimensions'] = result['updated_profile']
            save_profile(db, db_session.id, profile_data)
        
        save_record(db, db_session.id, {'type': 'evaluation', 'topic': '学习评估', 'detail': f"得分: {result.get('score', 0)}", 'score': result.get('score', 0)})
        
        return {"code": 0, "data": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/mistakes")
async def get_mistakes_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        mistakes = get_mistakes(db, db_session.id)
        
        total = len(mistakes)
        correct = 0
        wrong = total
        
        return {"code": 0, "data": {
            "mistakes": mistakes,
            "stats": {
                "total": total,
                "correct": correct,
                "wrong": wrong,
                "rate": round((correct / total) * 100) if total > 0 else 0
            }
        }}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/mistakes/{mistake_id}/review")
async def review_mistake_api(mistake_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        success = mark_mistake_reviewed(db, mistake_id)
        if success:
            return {"code": 0, "data": {"message": "标记已掌握成功"}}
        else:
            raise HTTPException(status_code=404, detail="错题不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/mistakes/{mistake_id}/unreview")
async def unreview_mistake_api(mistake_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        success = mark_mistake_unreviewed(db, mistake_id)
        if success:
            return {"code": 0, "data": {"message": "标记未掌握成功"}}
        else:
            raise HTTPException(status_code=404, detail="错题不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/agent/mistakes/{mistake_id}")
async def delete_mistake_api(mistake_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        success = delete_mistake(db, mistake_id)
        if success:
            return {"code": 0, "data": {"message": "错题删除成功"}}
        else:
            raise HTTPException(status_code=404, detail="错题不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/agent/mistakes")
async def clear_all_mistakes_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        clear_mistakes(db, db_session.id)
        return {"code": 0, "data": {"message": "全部错题已清空"}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/mistakes")
async def add_mistake_manual(mistake_data: dict = Body(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        save_mistake_full(db, db_session.id, {
            'question_id': mistake_data.get('question_id', 0),
            'question': mistake_data.get('question', ''),
            'question_type': mistake_data.get('question_type', 'choice'),
            'user_answer': mistake_data.get('user_answer', ''),
            'correct_answer': mistake_data.get('correct_answer', ''),
            'analysis': mistake_data.get('analysis', ''),
            'error_tags': mistake_data.get('error_tags', []),
            'error_detail': mistake_data.get('error_detail', ''),
            'related_knowledge': mistake_data.get('related_knowledge', ''),
            'subject': mistake_data.get('subject', '计算机'),
            'topic': mistake_data.get('topic', ''),
            'difficulty': mistake_data.get('difficulty', 'medium'),
            'related_document_id': mistake_data.get('related_document_id', 0),
            'source_type': mistake_data.get('source_type', 'manual')
        })
        
        return {"code": 0, "data": {"message": "错题录入成功"}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/mistakes/full")
async def get_mistakes_full_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        mistakes = get_mistakes_full(db, db_session.id)
        return {"code": 0, "data": mistakes}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/mistakes/subject/{subject}")
async def get_mistakes_by_subject_api(subject: str, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        mistakes = get_mistakes_by_subject(db, db_session.id, subject)
        return {"code": 0, "data": mistakes}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/mistakes/stats")
async def get_mistakes_stats_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        stats = get_mistakes_stats(db, db_session.id)
        return {"code": 0, "data": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/mistakes/trend")
async def get_mistakes_trend_api(days: int = 7, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        trend = get_daily_trend(db, db_session.id, days)
        return {"code": 0, "data": trend}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/mistakes/import")
async def import_mistakes(mistakes_data: List[Dict] = Body(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        for mistake in mistakes_data:
            save_mistake_full(db, db_session.id, {
                'question_id': mistake.get('question_id', 0),
                'question': mistake.get('question', ''),
                'question_type': mistake.get('question_type', 'choice'),
                'user_answer': mistake.get('user_answer', ''),
                'correct_answer': mistake.get('correct_answer', ''),
                'analysis': mistake.get('analysis', ''),
                'error_tags': mistake.get('error_tags', []),
                'error_detail': mistake.get('error_detail', ''),
                'related_knowledge': mistake.get('related_knowledge', ''),
                'subject': mistake.get('subject', '计算机'),
                'topic': mistake.get('topic', ''),
                'difficulty': mistake.get('difficulty', 'medium'),
                'related_document_id': mistake.get('related_document_id', 0),
                'source_type': 'import'
            })
        
        return {"code": 0, "data": {"message": f"成功导入 {len(mistakes_data)} 条错题"}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/mistakes/export")
async def export_mistakes(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        mistakes = get_mistakes_full(db, db_session.id)
        
        content = json.dumps(mistakes, ensure_ascii=False, indent=2)
        filename = f'mistakes_{datetime.now().strftime("%Y%m%d")}.json'
        
        export_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'exports')
        os.makedirs(export_dir, exist_ok=True)
        
        file_path = os.path.join(export_dir, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return FileResponse(file_path, filename=filename, media_type='application/json')
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/learning-loop")
async def trigger_learning_loop(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        mistakes = get_mistakes_full(db, db_session.id)
        profile = get_profile(db, db_session.id)
        plan = get_plan(db, db_session.id)
        
        if not profile:
            return {"code": 1, "data": {"message": "暂无学习画像，请先创建"}}
        
        profile_data = profile.profile_json
        
        results = automated_learning_loop(mistakes, profile_data, plan)
        
        if results.get('profile_updated'):
            save_profile(db, db_session.id, results['updated_profile'])
            print(f"Profile updated automatically for session {db_session.id}")
        
        if results.get('plan_adjusted'):
            save_plan(db, db_session.id, results['updated_plan'])
            print(f"Plan adjusted automatically for session {db_session.id}")
        
        return {
            "code": 0, 
            "data": {
                "message": "学习闭环自动化完成",
                "profile_updated": results.get('profile_updated', False),
                "plan_adjusted": results.get('plan_adjusted', False),
                "agent_status": get_agent_status()
            }
        }
    except Exception as e:
        print(f"Learning loop error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/records")
async def add_record(record_data: dict = Body(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        save_record(db, db_session.id, record_data)
        return {"code": 0, "data": record_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/records")
async def get_records_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        records = get_records(db, db_session.id)
        return {"code": 0, "data": records}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/status")
async def get_status():
    try:
        status = get_agent_status()
        return {"code": 0, "data": status}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/recommendations")
async def get_recommendations_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        profile = get_profile(db, db_session.id)
        plan = get_plan(db, db_session.id)
        
        if not profile:
            raise HTTPException(status_code=400, detail="请先创建学习画像")
        
        rec_list = generate_recommendations(profile.profile_json, plan)
        return {"code": 0, "data": {"recommendations": rec_list}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/loop")
async def run_loop(input_text: str = Body(..., embed=True), answers: Optional[List[Dict]] = Body(default=None), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        results = complete_learning_loop(input_text, answers)
        
        save_profile(db, db_session.id, results['profile'])
        save_plan(db, db_session.id, results['plan'])
        
        for key, content in results['resources'].items():
            save_resource(db, db_session.id, {'type': key, 'title': input_text, 'content': str(content), 'topic': input_text})
        
        if answers:
            for answer in answers:
                if answer.get('user_answer') != answer.get('correct_answer'):
                    save_mistake(db, db_session.id, {
                        'question': answer.get('question', ''),
                        'user_answer': answer.get('user_answer', ''),
                        'correct_answer': answer.get('correct_answer', ''),
                        'analysis': answer.get('explanation', ''),
                        'error_tags': [],
                        'related_knowledge': ''
                    })
        
        return {"code": 0, "data": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/document/upload")
async def upload_document(file: UploadFile = File(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        content = await file.read()
        file_size = len(content)
        
        file_type = file.filename.split('.')[-1].lower()
        
        content_str = ""
        analysis_result = {"topics": [], "key_points": []}
        
        if file_type == 'pdf':
            try:
                from PyPDF2 import PdfReader
                import io
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    content_str += page.extract_text() or ""
            except Exception as e:
                print(f"PDF parse error: {e}")
                content_str = "无法解析PDF内容，请确保PDF文件可读取"
        
        elif file_type == 'docx':
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(content))
                for paragraph in doc.paragraphs:
                    content_str += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content_str += cell.text + "\t"
                        content_str += "\n"
            except Exception as e:
                print(f"DOCX parse error: {e}")
                content_str = "无法解析DOCX文档，请确保文件格式正确"
        
        elif file_type == 'doc':
            content_str = ""
            try:
                import win32com.client
                import pythoncom
                import io
                import tempfile
                import os
                
                pythoncom.CoInitialize()
                
                with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(tmp_path)
                content_str = doc.Content.Text
                doc.Close()
                word.Quit()
                
                os.unlink(tmp_path)
                pythoncom.CoUninitialize()
            except Exception as e:
                print(f"DOC parse error (win32com): {e}")
                try:
                    import textract
                    content_str = textract.process(content).decode('utf-8', errors='ignore')
                except Exception as e2:
                    print(f"DOC parse error (textract): {e2}")
                    try:
                        content_str = content.decode('utf-8', errors='ignore')[:5000]
                    except:
                        content_str = "无法解析DOC文档，请将文件转换为DOCX格式或TXT格式"
        
        elif file_type in ['txt', 'md', 'json']:
            try:
                content_str = content.decode('utf-8')
            except:
                try:
                    content_str = content.decode('gbk', errors='ignore')
                except:
                    content_str = str(content)[:5000]
        
        else:
            try:
                content_str = content.decode('utf-8', errors='ignore')[:5000]
            except:
                content_str = f"文件类型 {file_type} 暂不支持文本提取"
        
        if content_str.strip() and len(content_str) > 10:
            analysis_result = extract_topics_from_document(content_str)
        
        core_topics = analysis_result.get('core_topics', [])
        knowledge_points = analysis_result.get('knowledge_points', [])
        
        doc = save_document(db, db_session.id, {
            'filename': file.filename,
            'file_type': file_type,
            'file_size': file_size,
            'content': content_str[:65535],
            'extracted_topics': core_topics,
            'full_analysis': analysis_result,
            'status': 'processed' if len(core_topics) > 0 else 'pending',
            'topics_count': len(core_topics)
        })
        
        if knowledge_points:
            for point in knowledge_points:
                save_knowledge_point(db, {
                    'document_id': doc.id,
                    'session_id': db_session.id,
                    'subject': point.get('subject', ''),
                    'chapter': point.get('chapter', ''),
                    'name': point.get('name', ''),
                    'definition': point.get('definition', ''),
                    'difficulty': point.get('difficulty', 'medium'),
                    'is_key': point.get('is_key', False),
                    'is_exam_point': point.get('is_exam_point', False),
                    'error_tags': point.get('error_tags', []),
                    'prerequisites': point.get('prerequisites', []),
                    'structured_data': point
                })
            
            doc.topics_count = len(knowledge_points)
            db.commit()
        
        return {"code": 0, "data": {
            "filename": file.filename,
            "file_size": file_size,
            "analysis": analysis_result,
            "status": 'processed' if len(core_topics) > 0 else 'pending',
            "topics_count": len(knowledge_points) if knowledge_points else len(core_topics),
            "knowledge_points_count": len(knowledge_points)
        }}
    except Exception as e:
        print(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/documents")
async def get_documents_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        docs = get_documents(db, db_session.id)
        return {"code": 0, "data": docs}
    except Exception as e:
        print(f"Get documents error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/document/{doc_id}")
async def get_document(doc_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        content = get_document_content(db, doc_id)
        if not content:
            raise HTTPException(status_code=404, detail="文档不存在")
        return {"code": 0, "data": content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/document/{doc_id}/extract-topics")
async def extract_topics(doc_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        content = get_document_content(db, doc_id)
        if not content:
            raise HTTPException(status_code=404, detail="文档不存在")
        
        if not content.get('content') or len(content['content']) < 10:
            raise HTTPException(status_code=400, detail="文档内容不足，无法提取知识点")
        
        db_session = get_or_create_session(db, session_key)
        
        doc = db.query(UploadedDocument).filter(UploadedDocument.id == doc_id).first()
        if not doc:
            raise HTTPException(status_code=404, detail="文档不存在")
        
        doc.status = 'processing'
        db.commit()
        
        analysis_result = extract_topics_from_document(content['content'])
        core_topics = analysis_result.get('core_topics', [])
        knowledge_points = analysis_result.get('knowledge_points', [])
        
        delete_knowledge_points_by_document(db, doc_id)
        
        for point in knowledge_points:
            save_knowledge_point(db, {
                'document_id': doc_id,
                'session_id': db_session.id,
                'subject': point.get('subject', doc.subject or ''),
                'chapter': point.get('chapter', doc.chapter or ''),
                'name': point.get('name', ''),
                'definition': point.get('definition', ''),
                'difficulty': point.get('difficulty', 'medium'),
                'is_key': point.get('is_key', False),
                'is_exam_point': point.get('is_exam_point', False),
                'error_tags': point.get('error_tags', []),
                'prerequisites': point.get('prerequisites', []),
                'structured_data': point
            })
        
        doc.extracted_topics = core_topics
        doc.full_analysis = analysis_result
        doc.status = 'processed' if len(core_topics) > 0 else 'error'
        doc.topics_count = len(knowledge_points) if knowledge_points else len(core_topics)
        db.commit()
        
        return {"code": 0, "data": {
            "document_id": doc_id,
            "analysis": analysis_result,
            "topics_count": len(knowledge_points) if knowledge_points else len(core_topics),
            "status": doc.status
        }}
    except Exception as e:
        print(f"Extract topics error: {e}")
        import traceback
        traceback.print_exc()
        
        doc = db.query(UploadedDocument).filter(UploadedDocument.id == doc_id).first()
        if doc:
            doc.status = 'error'
            db.commit()
        
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/document/{doc_id}/knowledge-points")
async def get_knowledge_points_api(doc_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        points = get_knowledge_points_by_document(db, doc_id)
        return {"code": 0, "data": [{
            'id': p.id,
            'document_id': p.document_id,
            'session_id': p.session_id,
            'subject': p.subject,
            'chapter': p.chapter,
            'name': p.name,
            'definition': p.definition,
            'difficulty': p.difficulty,
            'is_key': p.is_key,
            'is_exam_point': p.is_exam_point,
            'error_tags': p.error_tags,
            'prerequisites': p.prerequisites,
            'structured_data': p.structured_data,
            'created_at': p.created_at.isoformat()
        } for p in points]}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/knowledge-point/{point_id}")
async def update_knowledge_point_api(point_id: int, point_data: dict = Body(...), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        if point_id == 0:
            document_id = point_data.get('document_id')
            if not document_id:
                raise HTTPException(status_code=400, detail="文档ID不能为空")
            
            point = save_knowledge_point(db, {
                'document_id': document_id,
                'session_id': db_session.id,
                'subject': point_data.get('subject', ''),
                'chapter': point_data.get('chapter', ''),
                'name': point_data.get('name', ''),
                'definition': point_data.get('definition', ''),
                'difficulty': point_data.get('difficulty', 'medium'),
                'is_key': point_data.get('is_key', False),
                'is_exam_point': point_data.get('is_exam_point', False),
                'error_tags': point_data.get('error_tags', []),
                'prerequisites': point_data.get('prerequisites', []),
                'structured_data': {}
            })
            
            doc = db.query(UploadedDocument).filter(UploadedDocument.id == document_id).first()
            if doc:
                doc.topics_count = len(get_knowledge_points_by_document(db, document_id))
                db.commit()
            
            return {"code": 0, "data": {"message": "知识点添加成功"}}
        else:
            point = update_knowledge_point(db, point_id, point_data)
            if point:
                return {"code": 0, "data": {"message": "知识点更新成功"}}
            else:
                raise HTTPException(status_code=404, detail="知识点不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/agent/knowledge-point/{point_id}")
async def delete_knowledge_point_api(point_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        point = db.query(KnowledgePoint).filter(KnowledgePoint.id == point_id).first()
        if not point:
            raise HTTPException(status_code=404, detail="知识点不存在")
        
        db.delete(point)
        db.commit()
        
        doc = db.query(UploadedDocument).filter(UploadedDocument.id == point.document_id).first()
        if doc:
            remaining_points = get_knowledge_points_by_document(db, point.document_id)
            doc.topics_count = len(remaining_points)
            db.commit()
        
        return {"code": 0, "data": {"message": "知识点删除成功"}}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/knowledge-point/detail")
async def get_knowledge_point_detail(point_name: str, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        from app.database import User, MistakeRecord, LearningRecord, QuestionBank
        
        db_session = get_or_create_session(db, session_key)
        
        user = db.query(User).filter(User.session_id == db_session.id).first()
        session_id_to_use = db_session.id
        
        if user and user.session_id:
            session_id_to_use = user.session_id
        
        knowledge_points = get_knowledge_points(db, session_id_to_use)
        point = next((p for p in knowledge_points if p.name == point_name), None)
        
        if not point:
            return {"code": 1, "message": "知识点不存在"}
        
        mistakes = db.query(MistakeRecord).filter(
            MistakeRecord.session_id == session_id_to_use,
            MistakeRecord.related_knowledge == point_name
        ).all()
        
        questions = db.query(QuestionBank).filter(
            QuestionBank.session_id == session_id_to_use,
            QuestionBank.topic == point_name
        ).all()
        
        records = db.query(LearningRecord).filter(
            LearningRecord.session_id == session_id_to_use,
            LearningRecord.topic == point_name
        ).all()
        
        total_questions = len([r for r in records if r.type == 'question'])
        correct_questions = len([r for r in records if r.type == 'question' and r.score and r.score >= 60])
        
        if total_questions + len(mistakes) > 0:
            mastery = round((correct_questions / (total_questions + len(mistakes))) * 100)
        else:
            mastery = 0
        
        mistake_details = []
        for m in mistakes:
            mistake_details.append({
                'id': m.id,
                'question': m.question,
                'question_type': m.question_type,
                'user_answer': m.user_answer,
                'correct_answer': m.correct_answer,
                'analysis': m.analysis,
                'error_tags': m.error_tags,
                'error_detail': m.error_detail,
                'reviewed': m.reviewed,
                'created_at': m.created_at.isoformat()
            })
        
        question_details = []
        for q in questions:
            question_details.append({
                'id': q.id,
                'question': q.question,
                'question_type': q.question_type,
                'options': q.options,
                'correct_answer': q.correct_answer,
                'analysis': q.analysis,
                'difficulty': q.difficulty
            })
        
        point_data = {
            'id': point.id or 0,
            'name': point.name or '',
            'chapter': point.chapter or '',
            'subject': point.subject or '',
            'definition': point.definition or '',
            'difficulty': point.difficulty or 'medium',
            'is_key': point.is_key or False,
            'is_exam_point': point.is_exam_point or False,
            'prerequisites': point.prerequisites or [],
            'error_tags': point.error_tags or [],
            'mastery': mastery,
            'total_questions': total_questions + len(mistakes),
            'correct_questions': correct_questions,
            'mistakes': mistake_details,
            'questions': question_details
        }
        
        return {"code": 0, "data": point_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/knowledge/stats")
async def get_knowledge_stats_api(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        stats = get_knowledge_stats(db, db_session.id)
        return {"code": 0, "data": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/agent/knowledge-graph")
async def get_knowledge_graph(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        from app.database import User, MistakeRecord, LearningRecord
        
        db_session = get_or_create_session(db, session_key)
        
        user = db.query(User).filter(User.session_id == db_session.id).first()
        session_id_to_use = db_session.id
        
        if user and user.session_id:
            session_id_to_use = user.session_id
        
        knowledge_points = get_knowledge_points(db, session_id_to_use)
        
        mistakes = db.query(MistakeRecord).filter(MistakeRecord.session_id == session_id_to_use).all()
        records = db.query(LearningRecord).filter(LearningRecord.session_id == session_id_to_use).all()
        
        point_mistake_counts = {}
        for m in mistakes:
            related_knowledge = m.related_knowledge or ''
            if related_knowledge:
                point_mistake_counts[related_knowledge] = point_mistake_counts.get(related_knowledge, 0) + 1
        
        point_question_counts = {}
        point_correct_counts = {}
        for r in records:
            if r.type == 'question' and r.topic:
                point_question_counts[r.topic] = point_question_counts.get(r.topic, 0) + 1
                if r.score and r.score >= 60:
                    point_correct_counts[r.topic] = point_correct_counts.get(r.topic, 0) + 1
        
        nodes = []
        links = []
        mastered_count = 0
        weak_count = 0
        
        for idx, point in enumerate(knowledge_points):
            point_name = point.name or ''
            
            total_questions = point_question_counts.get(point_name, 0)
            correct_questions = point_correct_counts.get(point_name, 0)
            mistake_count = point_mistake_counts.get(point_name, 0)
            
            mastery = None
            has_data = False
            
            if total_questions > 0:
                mastery = round((correct_questions / total_questions) * 100)
                has_data = True
            elif mistake_count > 0:
                mastery = max(0, 50 - mistake_count * 10)
                has_data = True
            
            if has_data and mastery >= 80:
                mastered_count += 1
            elif has_data and mastery < 40:
                weak_count += 1
            
            node = {
                'id': idx + 1,
                'name': point_name,
                'chapter': point.chapter or '',
                'difficulty': point.difficulty or 'medium',
                'mastery': mastery,
                'is_exam_point': point.is_exam_point or False,
                'definition': point.definition or '',
                'is_key': point.is_key or False,
                'subject': point.subject or '',
                'mistake_count': mistake_count,
                'total_questions': total_questions
            }
            nodes.append(node)
            
            prerequisites = point.prerequisites or []
            for prereq in prerequisites:
                prereq_node = next((n for n in nodes if n['name'] == prereq), None)
                if prereq_node:
                    links.append({
                        'source': prereq_node['id'],
                        'target': idx + 1
                    })
        
        stats = {
            'total': len(nodes),
            'relations': len(links),
            'mastered': mastered_count,
            'weak': weak_count
        }
        
        recommendations = []
        if weak_count > 0:
            recommendations.append({
                'icon': '📚',
                'title': '重点复习薄弱知识点',
                'description': f'你有 {weak_count} 个知识点掌握度较低，建议优先复习'
            })
        if mastered_count > 0:
            recommendations.append({
                'icon': '🎯',
                'title': '挑战更高难度',
                'description': f'你已掌握 {mastered_count} 个知识点，可以尝试更难的内容'
            })
        recommendations.append({
            'icon': '💻',
            'title': '多做编程练习',
            'description': '通过实践加深对知识点的理解和应用能力'
        })
        
        return {
            'code': 0,
            'data': {
                'nodes': nodes,
                'links': links,
                'stats': stats,
                'recommendations': recommendations
            }
        }
    
    except Exception as e:
        print(f"Knowledge graph error: {e}")
        import traceback
        traceback.print_exc()
        return {
            'code': 0,
            'data': {
                'nodes': [],
                'links': [],
                'stats': {'total': 0, 'relations': 0, 'mastered': 0, 'weak': 0},
                'recommendations': []
            }
        }

@router.delete("/agent/graph/clear")
async def clear_graph(session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        clear_all_knowledge_points(db, db_session.id)
        return {'code': 0, 'message': '图谱已清空'}
    except Exception as e:
        print(f"Clear graph error: {e}")
        return {'code': -1, 'message': str(e)}

@router.delete("/agent/document/{doc_id}")
async def delete_document_endpoint(doc_id: int, session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        success = delete_document(db, doc_id)
        if success:
            return {"code": 0, "data": {"message": "文档删除成功"}}
        else:
            raise HTTPException(status_code=404, detail="文档不存在")
    except Exception as e:
        print(f"Delete document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/ocr")
async def ocr_analyze(image_text: str = Body(..., embed=True), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        result = ocr_analyze_image(image_text)
        
        if result.get('has_questions'):
            for q in result.get('questions', []):
                save_mistake(db, db_session.id, {
                    'question': q.get('question', ''),
                    'user_answer': '',
                    'correct_answer': q.get('answer', ''),
                    'analysis': q.get('explanation', ''),
                    'error_tags': [],
                    'related_knowledge': result.get('subject', '')
                })
        
        return {"code": 0, "data": result}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/ocr/image")
async def ocr_analyze_image_api(file: UploadFile = File(...), session_key: Optional[str] = Cookie(None)):
    try:
        import tempfile
        import os
        
        contents = await file.read()
        print(f"Uploaded file size: {len(contents)} bytes")
        print(f"Uploaded file type: {file.content_type}")
        
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            f.write(contents)
            temp_path = f.name
        
        print(f"Temp file path: {temp_path}")
        
        prompt = "识别这张图片中的所有文字内容，只输出识别结果，不要任何其他解释。"
        
        try:
            import dashscope
            from dashscope import Generation
            dashscope.api_key = os.getenv("DASHSCOPE_API_KEY", "")
            
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image", "image": temp_path}
                    ]
                }
            ]
            
            print(f"Calling qwen-vl-plus model...")
            response = Generation.call(
                model="qwen-vl-plus",
                messages=messages,
                result_format="message"
            )
            
            print(f"OCR Response: {response}")
            
            extracted_text = response.output.choices[0].message.content
            print(f"Extracted text: {extracted_text[:500]}")
            
            return {"code": 0, "data": {"text": extracted_text}}
            
        except Exception as e:
            print(f"Generation error: {e}")
            import traceback
            traceback.print_exc()
            
            try:
                from dashscope import MultiModalConversation
                
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image", "image": temp_path}
                        ]
                    }
                ]
                
                print(f"Calling MultiModalConversation...")
                response = MultiModalConversation.call(
                    model="qwen-vl-plus",
                    messages=messages
                )
                
                print(f"MultiModal Response: {response}")
                
                choices = response.output.choices
                if choices and len(choices) > 0:
                    message = choices[0].message
                    if message and message.content:
                        content_list = message.content
                        if isinstance(content_list, list) and len(content_list) > 0:
                            first_item = content_list[0]
                            if isinstance(first_item, dict) and 'text' in first_item:
                                extracted_text = first_item['text']
                            elif hasattr(first_item, 'text'):
                                extracted_text = first_item.text
                            else:
                                extracted_text = str(first_item)
                        elif isinstance(content_list, str):
                            extracted_text = content_list
                        else:
                            extracted_text = str(content_list)
                    else:
                        extracted_text = str(message) if message else ""
                else:
                    extracted_text = ""
                
                print(f"Extracted text: {extracted_text[:500]}")
                return {"code": 0, "data": {"text": extracted_text}}
                
            except Exception as e2:
                print(f"MultiModal error: {e2}")
                traceback.print_exc()
                return {"code": 1, "data": {"text": "", "error": f"AI识别失败: {str(e2)}"}}
        finally:
            os.unlink(temp_path)
            
    except Exception as e:
        print(f"OCR API error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/agent/export")
async def export_data(data_type: str = Body(..., embed=True), session_key: Optional[str] = Cookie(None), db: Session = Depends(get_db)):
    try:
        db_session = get_or_create_session(db, session_key)
        
        if data_type == 'profile':
            profile = get_profile(db, db_session.id)
            content = json.dumps(profile.profile_json if profile else {}, ensure_ascii=False, indent=2)
            filename = 'profile.json'
        elif data_type == 'plan':
            plan = get_plan(db, db_session.id)
            content = json.dumps(plan, ensure_ascii=False, indent=2)
            filename = 'plan.json'
        elif data_type == 'mistakes':
            mistakes = get_mistakes(db, db_session.id)
            content = json.dumps(mistakes, ensure_ascii=False, indent=2)
            filename = 'mistakes.json'
        elif data_type == 'records':
            records = get_records(db, db_session.id)
            content = json.dumps(records, ensure_ascii=False, indent=2)
            filename = 'records.json'
        else:
            raise HTTPException(status_code=400, detail="不支持的导出类型")
        
        export_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'exports')
        os.makedirs(export_dir, exist_ok=True)
        
        file_path = os.path.join(export_dir, filename)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return FileResponse(file_path, filename=filename, media_type='application/json')
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
